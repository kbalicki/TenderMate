import asyncio
import json
import logging
import os
import re
from pathlib import Path

from app.config import settings

logger = logging.getLogger(__name__)

# Supported binary file extensions and their text extractors
_BINARY_EXTENSIONS = {'.docx', '.doc', '.pdf', '.xlsx', '.xls'}


def _extract_text_from_file(file_path: Path) -> str | None:
    """Extract text content from a file. Returns None if extraction fails."""
    suffix = file_path.suffix.lower()

    if suffix == '.docx':
        try:
            import docx
            doc = docx.Document(str(file_path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        paragraphs.append(" | ".join(cells))
            text = "\n".join(paragraphs)
            logger.info(f"[FileExtract] DOCX {file_path.name}: {len(text)} znaków, {len(paragraphs)} akapitów")
            return text
        except Exception as e:
            logger.warning(f"[FileExtract] Błąd DOCX {file_path.name}: {e}")
            return None

    elif suffix == '.pdf':
        try:
            import fitz  # pymupdf
            doc = fitz.open(str(file_path))
            pages = []
            for page in doc:
                pages.append(page.get_text())
            doc.close()
            text = "\n\n".join(pages)
            logger.info(f"[FileExtract] PDF {file_path.name}: {len(text)} znaków, {len(pages)} stron")
            return text
        except Exception as e:
            logger.warning(f"[FileExtract] Błąd PDF {file_path.name}: {e}")
            return None

    elif suffix == '.xlsx':
        try:
            import openpyxl
            wb = openpyxl.load_workbook(str(file_path), read_only=True)
            lines = []
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                lines.append(f"--- Arkusz: {sheet} ---")
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c) if c is not None else "" for c in row]
                    if any(cells):
                        lines.append(" | ".join(cells))
            wb.close()
            text = "\n".join(lines)
            logger.info(f"[FileExtract] XLSX {file_path.name}: {len(text)} znaków")
            return text
        except Exception as e:
            logger.warning(f"[FileExtract] Błąd XLSX {file_path.name}: {e}")
            return None

    else:
        # Plain text file
        try:
            text = file_path.read_text(errors="ignore")
            return text
        except Exception as e:
            logger.warning(f"[FileExtract] Błąd odczytu {file_path.name}: {e}")
            return None


def _fix_json_quotes(text: str) -> str:
    """Fix unescaped quotes inside JSON string values by iterative repair.

    Claude sometimes produces JSON where string values contain unescaped
    double quotes (e.g. Polish „..." or citation marks). This function
    replaces typographic quotes and then iteratively fixes parse errors
    by escaping the problematic quote at each error position.
    """
    # Replace Polish typographic quotes with single quotes (decorative, not structural)
    text = text.replace('\u201e', "'")  # „
    text = text.replace('\u201d', "'")  # "
    text = text.replace('\u201c', "'")  # "
    text = text.replace('\u00ab', "'")  # «
    text = text.replace('\u00bb', "'")  # »

    # Iteratively fix unescaped quotes: find parse error, escape the quote, retry
    for _ in range(50):  # Max 50 fixes
        try:
            json.loads(text)
            return text  # Valid JSON
        except json.JSONDecodeError as e:
            # Find the problematic position and look for a quote nearby to escape
            pos = e.pos
            # The error is usually at a " that the parser interprets as end-of-string
            # but it's actually inside the string. Look backwards for the quote.
            # Strategy: look at chars around error pos for unescaped "
            search_start = max(0, pos - 3)
            search_end = min(len(text), pos + 3)
            fixed = False
            for check_pos in range(search_start, search_end):
                if text[check_pos] == '"' and (check_pos == 0 or text[check_pos - 1] != '\\'):
                    # Check if escaping this quote would help
                    candidate = text[:check_pos] + '\\"' + text[check_pos + 1:]
                    try:
                        json.loads(candidate)
                        return candidate
                    except json.JSONDecodeError:
                        text = candidate
                        fixed = True
                        break
            if not fixed:
                break  # Can't fix
    return text


def _extract_json(text: str) -> dict | list | None:
    """Try to extract JSON from text that may contain markdown fences or other wrapping."""
    # Try direct parse first
    try:
        return json.loads(text)
    except (json.JSONDecodeError, TypeError):
        pass

    if not isinstance(text, str):
        return None

    # Strip markdown code fences: ```json ... ``` or ``` ... ```
    match = re.search(r'```(?:json)?\s*\n?(.*?)\n?\s*```', text, re.DOTALL)
    if match:
        inner = match.group(1).strip()
        try:
            return json.loads(inner)
        except json.JSONDecodeError:
            # Try fixing typographic quotes
            try:
                return json.loads(_fix_json_quotes(inner))
            except json.JSONDecodeError:
                pass

    # Try to find first { ... } or [ ... ] block
    for start_char, end_char in [('{', '}'), ('[', ']')]:
        start = text.find(start_char)
        if start == -1:
            continue
        # Find matching closing bracket from the end
        end = text.rfind(end_char)
        if end > start:
            candidate = text[start:end + 1]
            try:
                return json.loads(candidate)
            except json.JSONDecodeError:
                # Try fixing typographic quotes
                try:
                    return json.loads(_fix_json_quotes(candidate))
                except json.JSONDecodeError:
                    pass

    return None


async def call_claude(
    prompt: str,
    context_files: list[Path] | None = None,
    system_prompt: str | None = None,
    json_schema: dict | None = None,
) -> dict | str:
    """Call Claude CLI via subprocess in print mode."""
    cmd = [
        "claude",
        "-p",
        "--model", settings.claude_model,
        "--output-format", "json",
    ]

    if system_prompt:
        cmd.extend(["--system-prompt", system_prompt])

    if json_schema:
        cmd.extend(["--json-schema", json.dumps(json_schema)])

    # Build prompt with file contents inlined
    full_prompt = prompt
    if context_files:
        for f in context_files:
            if not f.exists():
                continue
            if f.stat().st_size > 2_000_000:
                logger.warning(f"[Claude CLI] Plik {f.name} za duży ({f.stat().st_size} B), pomijam")
                continue

            content = _extract_text_from_file(f)
            if content and len(content.strip()) > 10:
                # Truncate very long files to avoid prompt overflow
                if len(content) > 50_000:
                    content = content[:50_000] + "\n\n[... tekst obcięty, oryginał ma " + str(len(content)) + " znaków]"
                full_prompt += f"\n\n--- PLIK: {f.name} ---\n{content}\n"
            else:
                logger.warning(f"[Claude CLI] Nie udało się wyekstrahować tekstu z {f.name}")

    prompt_preview = full_prompt[:300].replace("\n", " ")
    logger.info(f"[Claude CLI] Wywołuję (model={settings.claude_model}), prompt: {prompt_preview}...")
    logger.info(f"[Claude CLI] Pliki kontekstowe: {[f.name for f in (context_files or []) if f.exists()]}")
    logger.info(f"[Claude CLI] Długość promptu: {len(full_prompt)} znaków")

    # Remove CLAUDECODE env var to avoid "nested session" error
    env = {k: v for k, v in os.environ.items() if k != "CLAUDECODE"}

    proc = await asyncio.create_subprocess_exec(
        *cmd,
        stdin=asyncio.subprocess.PIPE,
        stdout=asyncio.subprocess.PIPE,
        stderr=asyncio.subprocess.PIPE,
        env=env,
    )

    stdout, stderr = await proc.communicate(input=full_prompt.encode())

    stderr_text = stderr.decode()
    if stderr_text:
        logger.warning(f"[Claude CLI] stderr: {stderr_text[:500]}")

    if proc.returncode != 0:
        err_msg = f"Claude CLI błąd (exit {proc.returncode}): {stderr_text}"
        logger.error(f"[Claude CLI] {err_msg}")
        raise RuntimeError(err_msg)

    result_text = stdout.decode()
    logger.info(f"[Claude CLI] Odpowiedź: {len(result_text)} znaków")
    logger.debug(f"[Claude CLI] Surowa odpowiedź (pierwsze 500 zn.): {result_text[:500]}")

    # Parse the Claude CLI JSON envelope
    parsed = _extract_json(result_text)
    if parsed is None:
        logger.warning(f"[Claude CLI] Nie udało się sparsować JSON, zwracam surowy tekst: {result_text[:200]}")
        return result_text

    # Claude CLI --output-format json wraps result: {"type":"result","result":"..."}
    if isinstance(parsed, dict) and "result" in parsed and "type" in parsed:
        inner = parsed["result"]
        logger.info(f"[Claude CLI] Rozpakowano envelope, typ inner: {type(inner).__name__}")

        if isinstance(inner, str):
            extracted = _extract_json(inner)
            if extracted is not None:
                logger.info(f"[Claude CLI] Sparsowano JSON z inner string, klucze: {list(extracted.keys()) if isinstance(extracted, dict) else 'lista'}")
                return extracted
            logger.warning(f"[Claude CLI] Inner string nie jest JSON: {inner[:200]}")
            return inner

        if isinstance(inner, (dict, list)):
            logger.info(f"[Claude CLI] Inner jest już dict/list, klucze: {list(inner.keys()) if isinstance(inner, dict) else 'lista'}")
            return inner

    logger.info(f"[Claude CLI] Zwracam parsed bez envelope, typ: {type(parsed).__name__}")
    return parsed
